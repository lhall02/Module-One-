from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs")
AUTHOR = "Logan Hall"


def set_font(run, name="Arial", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_doc(doc, title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title_style = styles["Title"]
    title_style.font.name = "Arial"
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(10)

    for style_name, size, before, after in [
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    doc.core_properties.author = AUTHOR
    doc.core_properties.title = title


def paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run)
    return p


def title(doc, text, subtitle=None):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=24, bold=True)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(14)
        run2 = p2.add_run(subtitle)
        set_font(run2, size=11, color="555555", italic=True)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_font(run, size={1: 16, 2: 14, 3: 12}[level], bold=True)
    return p


def save_doc(filename, title_text, subtitle, sections):
    doc = Document()
    configure_doc(doc, title_text)
    title(doc, title_text, subtitle)
    for heading_text, paragraphs in sections:
        heading(doc, heading_text, 1)
        for text in paragraphs:
            paragraph(doc, text)
    doc.save(OUT_DIR / filename)


def build():
    OUT_DIR.mkdir(exist_ok=True)

    save_doc(
        "professional-self-assessment.docx",
        "Professional Self-Assessment",
        "CS 499 Computer Science Capstone | Logan Hall",
        [
            (
                "Program Growth and Professional Direction",
                [
                    "Throughout the Computer Science program, I moved from learning individual programming concepts to thinking about complete systems. Courses in programming, databases, algorithms, testing, full-stack development, and security helped me understand that professional software is more than code that runs. It also needs to be readable, maintainable, secure, testable, and useful to the people who depend on it.",
                    "Completing this ePortfolio helped me showcase that growth because it required me to revisit an existing project, evaluate its weaknesses, and improve it in several connected areas. The Travlr application gave me a practical way to demonstrate full-stack development skills through customer-facing pages, Express routes and controllers, API endpoints, MongoDB/Mongoose modeling, and an Angular administration interface.",
                    "My professional goal is to continue growing toward software engineering and full-stack web development roles. I want to build applications that solve real problems, communicate clearly with users and teammates, and handle data responsibly.",
                ],
            ),
            (
                "Collaboration, Communication, and Stakeholders",
                [
                    "The capstone code review strengthened my ability to communicate about software in a structured way. I practiced describing current behavior, identifying risks, proposing improvements, and connecting technical work to stakeholder needs. Even when working individually, I treated the project as a shared codebase that another developer, instructor, or manager would need to understand later.",
                    "The ePortfolio itself also demonstrates communication because it presents source code, narratives, and reflection in a format designed for an outside audience. The goal is not only to show that the code exists, but also to explain why the enhancements matter.",
                ],
            ),
            (
                "Technical Strengths",
                [
                    "The software engineering work demonstrates MVC structure, route planning, view rendering, RESTful API design, and maintainability. The database work demonstrates Mongoose schema design, validation, unique identifiers, API error handling, and Angular service integration.",
                    "The algorithms and data structures work demonstrates binary search, filtering, grouping, date-range logic, and Map-based caching. These features show that I can apply algorithmic thinking to real application behavior rather than only to isolated classroom examples.",
                    "The security mindset developed through this work focuses on validation, defensive error handling, database integrity, and reducing unsafe query behavior. Security is not a single feature added at the end; it is a habit that shapes each route, model, and data flow.",
                ],
            ),
            (
                "How the Artifacts Fit Together",
                [
                    "The three enhanced artifacts work together as one story of improving a full-stack application. The software design artifact improves the customer workflow and project organization. The algorithms and data structures artifact improves how trip data is searched, filtered, grouped, cached, and retrieved. The database artifact improves data quality, API behavior, and administrative management.",
                    "Together, these artifacts demonstrate that I can evaluate an existing system, identify meaningful improvements, implement changes that support users and future developers, and communicate technical decisions in a professional portfolio.",
                ],
            ),
        ],
    )

    save_doc(
        "code-review-journal.docx",
        "Informal Code Review Journal",
        "CS 499 Computer Science Capstone | Logan Hall",
        [
            (
                "Purpose of Code Review",
                [
                    "Code review is the process of examining existing source code to understand what it does, identify weaknesses, and suggest improvements before the code is enhanced, merged, or released. It is important for computer science professionals because it improves code quality, supports teamwork, catches bugs earlier, and helps developers create software that is easier to maintain, update, and secure.",
                    "Best practices used for this capstone include reviewing small sections of code at a time, using a checklist, testing whether code works as intended, checking naming and formatting, looking for repeated or unreachable code, reviewing comments and documentation, and identifying security concerns such as weak validation or poor error handling.",
                ],
            ),
            (
                "Recording Approach",
                [
                    "For the code review recording, I selected ScreenPal because it can record, edit, and save a video while I walk through the code. My approach is to prepare a script before recording so the review stays organized and addresses the rubric areas clearly.",
                    "The script is organized around software engineering and design, algorithms and data structures, and databases. For each category, I explain what the current code does, analyze weaknesses, and describe the planned enhancement.",
                ],
            ),
            (
                "Enhancement Focus",
                [
                    "For software engineering and design, the review focuses on structure, readability, documentation, error handling, and maintainability. The completed enhancement adds a dedicated trip detail workflow and clearer MVC routing.",
                    "For algorithms and data structures, the review focuses on trip search, sorting, filtering, grouping, date-range logic, and caching. These features demonstrate algorithmic reasoning and data structure trade-offs.",
                    "For databases, the review focuses on how data is stored, retrieved, validated, and protected. The completed enhancement strengthens the Mongoose schema, API validation, CRUD behavior, and Angular admin integration.",
                ],
            ),
        ],
    )

    save_doc(
        "software-design-engineering-narrative.docx",
        "Software Design and Engineering Narrative",
        "Artifact One | Travlr Getaways",
        [
            (
                "Artifact Description",
                [
                    "The selected artifact is the Travlr Getaways web application, created from coursework involving full-stack web development with Express, Handlebars, Angular, and MongoDB. The software design enhancement focuses on improving the customer-facing trip browsing workflow.",
                    "The original project displayed trip content, but it did not provide a dedicated trip detail route. The enhanced version allows a user to open the travel page, select a trip, and view a focused details page for that trip.",
                ],
            ),
            (
                "Enhancement Completed",
                [
                    "The enhancement adds a `/travel/:code` route, controller logic that retrieves one trip through the API, and a new Handlebars detail view. The travel list now links each trip image, title, and action button to the correct detail page.",
                    "The controller handles missing trip codes, not-found trips, API failures, and server errors. This improves the user experience while keeping the code organized across routes, controllers, views, and API endpoints.",
                ],
            ),
            (
                "Reflection and Course Outcomes",
                [
                    "This enhancement demonstrates MVC design, route planning, view rendering, user-centered workflow design, and maintainability. It improved the artifact by making the application easier to navigate and by making the code easier for another developer to understand.",
                    "The work supports course outcomes related to professional-quality communication, designing computing solutions using computer science practices, using well-founded tools and techniques, and supporting collaboration through clearer code organization.",
                ],
            ),
        ],
    )

    save_doc(
        "algorithms-data-structures-narrative.docx",
        "Algorithms and Data Structures Narrative",
        "Artifact Two | Travlr Getaways",
        [
            (
                "Artifact Description",
                [
                    "The selected artifact is the trip data portion of the Travlr application. It is a full-stack web application that uses a Node.js and Express backend with MongoDB for storing trip data. For this milestone, the focus is the API controller and routes that manage trip records.",
                    "I selected this artifact because the trip data portion of the application was a strong place to demonstrate algorithms and data structure skills in a practical way. A travel application needs to search, organize, filter, group, and retrieve data efficiently.",
                ],
            ),
            (
                "Enhancements Completed",
                [
                    "The original API included basic trip retrieval. The enhanced version adds binary search by trip code, a Map-based cache with a time-to-live value, multi-criteria filtering, grouping by resort, and date-range search.",
                    "Binary search demonstrates the difference between linear search and O(log n) lookup after data is sorted. The cache demonstrates a space-time trade-off by storing frequently requested trip data for a short period. Filtering, grouping, and date-range logic make the API more flexible for real application use.",
                ],
            ),
            (
                "Reflection and Course Outcomes",
                [
                    "Through this enhancement, I learned that algorithms and data structures are useful beyond classroom examples. The binary search, cache, filtering logic, grouping, and date range queries all solve realistic problems in the application.",
                    "The enhancement supports the algorithms and data structures outcome because it uses searching, filtering, grouping, caching, and range query logic to solve data access problems. It also supports the tools and techniques outcome through RESTful API design and MongoDB aggregation, and it supports the security outcome through input validation and error handling.",
                ],
            ),
        ],
    )

    save_doc(
        "databases-narrative.docx",
        "Databases Narrative",
        "Artifact Three | Travlr Getaways",
        [
            (
                "Artifact Description",
                [
                    "The selected database artifact is the MongoDB/Mongoose trip model and the Travlr API routes that create, retrieve, update, delete, and query trip records. The artifact originated from coursework that used MongoDB with a Node/Express application.",
                    "The project already had trip fields such as name, code, start date, resort, price, image, and description. The enhancement strengthens how those records are validated, queried, and managed through the API and Angular admin interface.",
                ],
            ),
            (
                "Enhancement Completed",
                [
                    "The Trip schema includes required fields, trimming, uppercase trip codes, a unique trip code rule, minimum price validation, and timestamps. The API validates object IDs, required fields, numeric price filters, date filters, and date ranges before completing database operations.",
                    "The API also supports CRUD operations, filtered retrieval, grouped resort statistics, date-range lookup, and cache clearing. The Angular admin service and CRUD components connect to `/api/trips` so trip records can be listed, added, edited, and deleted.",
                ],
            ),
            (
                "Reflection and Course Outcomes",
                [
                    "This enhancement reinforced the importance of thoughtful schema design and controlled data access. Data services are not just about storing information; they must support the application purpose safely, consistently, and effectively.",
                    "The database enhancement supports the course outcome related to security because it anticipates invalid input and protects data integrity. It also supports the outcomes for designing computing solutions and using well-founded tools because the work uses Mongoose, Express routing, RESTful conventions, and Angular services.",
                ],
            ),
        ],
    )


if __name__ == "__main__":
    build()
